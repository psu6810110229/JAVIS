"""actuators/document_tools.py - Document generation tools for Jarvis."""
from __future__ import annotations

import asyncio
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from app.actuators.registry import tool

_DEFAULT_REPORTS_DIR = Path.home() / "Documents" / "JarvisReports"
_INVALID_FILENAME = re.compile(r'[<>:"/\\|?*]+')


def _safe_filename(name: str) -> str:
    cleaned = _INVALID_FILENAME.sub("_", name.strip())
    cleaned = cleaned.strip(" .")
    return cleaned or "jarvis_report"


@tool(
    name="write_docx_report",
    description=(
        "Create a Microsoft Word .docx report file with optional title and heading sections. "
        "Use this when the user asks for a report/document in Word format."
    ),
    parameters={
        "type": "object",
        "properties": {
            "title": {
                "type": "string",
                "description": "Document title shown at the top of the report.",
            },
            "content": {
                "type": "string",
                "description": "Main report body text. Use newline breaks for paragraphs.",
            },
            "file_name": {
                "type": "string",
                "description": "Optional output filename without extension. Defaults to report with timestamp.",
            },
            "output_dir": {
                "type": "string",
                "description": "Optional absolute output directory path. Defaults to Documents/JarvisReports.",
            },
        },
        "required": ["content"],
    },
    risk_level="low",
    category="documents",
    timeout_seconds=20,
)
async def write_docx_report(
    content: str,
    title: str = "",
    file_name: str = "",
    output_dir: str = "",
) -> dict[str, Any]:
    """Create a .docx report file and return file metadata."""
    if not content.strip():
        return {
            "status": "failed",
            "verified": False,
            "evidence": "No content provided for report generation.",
            "error": "content cannot be empty.",
        }

    target_dir = Path(output_dir).expanduser() if output_dir.strip() else _DEFAULT_REPORTS_DIR
    if not target_dir.is_absolute():
        return {
            "status": "failed",
            "verified": False,
            "evidence": "Output directory must be an absolute path.",
            "error": "output_dir must be absolute.",
        }

    chosen_name = file_name.strip()
    if not chosen_name:
        chosen_name = datetime.now().strftime("jarvis_report_%Y%m%d_%H%M%S")
    stem = _safe_filename(chosen_name)
    target_path = target_dir / f"{stem}.docx"

    def _write_docx() -> None:
        from docx import Document  # type: ignore[import]

        target_dir.mkdir(parents=True, exist_ok=True)
        document = Document()
        if title.strip():
            document.add_heading(title.strip(), level=1)

        for raw_line in content.splitlines():
            line = raw_line.strip()
            if not line:
                document.add_paragraph("")
                continue
            document.add_paragraph(line)

        document.save(str(target_path))

    try:
        await asyncio.to_thread(_write_docx)
    except ImportError as err:
        return {
            "status": "failed",
            "verified": False,
            "evidence": "python-docx dependency is missing.",
            "error": str(err),
        }
    except OSError as err:
        return {
            "status": "failed",
            "verified": False,
            "evidence": "Operating system failed to write the report file.",
            "error": str(err),
        }

    return {
        "status": "ok",
        "verified": target_path.exists(),
        "evidence": f"Report created at {target_path}",
        "file_path": str(target_path),
        "title": title.strip() or None,
        "bytes": target_path.stat().st_size if target_path.exists() else 0,
    }
